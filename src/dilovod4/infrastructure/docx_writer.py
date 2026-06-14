"""DocxDocumentWriter — адаптер порту DocumentWriter на базі python-docx.

Фізично відтворює оформлення згідно з параметрами доменного Document:
поля (§6.2), гарнітура та кеглі (§7.2), міжрядковий інтервал (§7.3),
відступи (§7.7), нумерація сторінок (§7.10). Наповнює текстом з DocumentContent.

Це інфраструктура: уся залежність від python-docx ізольована тут.
"""

from __future__ import annotations

import io

from docx import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from ..domain.model import Document, DocumentContent
from ..domain.model.qr_payload import build_signature_qr_payload

# §5.10: QR-код — рівно 21×21 мм.
_QR_SIDE_MM = 21
# §7.2: довідкові дані 8–12 pt — беремо мінімум для компактності відмітки.
_MARK_FONT_PT = 8


class DocxDocumentWriter:
    """Записує доменний документ у файл .docx за правилами ДСТУ 4163:2020."""

    def write(self, document: Document, content: DocumentContent, destination: str) -> str:
        doc = DocxDocument()
        self._apply_page_geometry(doc, document)
        self._apply_base_style(doc, document)
        self._enable_page_numbering(doc, document)

        self._add_letterhead(doc, document, content)
        self._add_approval(doc, document, content)
        self._add_addressees(doc, document, content)
        self._add_title(doc, document, content)
        self._add_body(doc, document, content)
        self._add_signature(doc, document, content)
        self._add_agreements_and_visas(doc, document, content)

        if not destination.endswith(".docx"):
            destination = f"{destination}.docx"
        doc.save(destination)
        return destination

    # --- §6.2 поля ---
    def _apply_page_geometry(self, doc, document: Document) -> None:
        m = document.geometry.margins
        for section in doc.sections:
            section.left_margin = Mm(m.left)
            section.right_margin = Mm(m.right)
            section.top_margin = Mm(m.top)
            section.bottom_margin = Mm(m.bottom)
            # §6.1 формат: A4 210x297, A5 210x148, A3 297x420
            dims = {
                "A4": (210, 297),
                "A5": (210, 148),
                "A3": (297, 420),
            }[document.geometry.paper_format.value]
            section.page_width = Mm(dims[0])
            section.page_height = Mm(dims[1])

    # --- §7.2 гарнітура/кегль, §7.3 інтервал ---
    def _apply_base_style(self, doc, document: Document) -> None:
        style = doc.styles["Normal"]
        font = style.font
        if document.typography.is_times_new_roman:
            font.name = "Times New Roman"
            # кирилиця: явно вказуємо east-asia/cs шрифт через rPr
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:ascii"), "Times New Roman")
            rfonts.set(qn("w:hAnsi"), "Times New Roman")
            rfonts.set(qn("w:cs"), "Times New Roman")
        font.size = Pt(document.typography.body_size_pt)

        pf = style.paragraph_format
        spacing = document.line_spacing.body_spacing
        pf.line_spacing = spacing
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.space_after = Pt(0)

    # --- §7.10 нумерація сторінок ---
    def _enable_page_numbering(self, doc, document: Document) -> None:
        section = doc.sections[0]
        # перша сторінка не нумерується (§7.10)
        section.different_first_page_header_footer = True

        # верхній колонтитул решти сторінок: посередині, арабські, без слова
        header = section.header
        header.is_linked_to_previous = False
        para = header.paragraphs[0]
        para.text = ""
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._insert_page_field(para)

        # перша сторінка — порожній колонтитул
        first = section.first_page_header
        first.is_linked_to_previous = False
        if first.paragraphs:
            first.paragraphs[0].text = ""

    def _insert_page_field(self, paragraph) -> None:
        """Вставити поле PAGE (арабська цифра), без тексту «сторінка»."""
        run = paragraph.add_run()
        fld_begin = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        instr = run._r.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
        instr.text = "PAGE \\* ARABIC"
        fld_end = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    # --- реквізити бланка ---
    def _add_letterhead(self, doc, document: Document, content: DocumentContent) -> None:
        # робоча позначка (напр. «ПРОЕКТ») — праворуч угорі, відступ 100 мм
        if content.marking.strip():
            mk = doc.add_paragraph()
            mk.paragraph_format.left_indent = Mm(document.left_indents.approval_mm)
            mk.add_run(content.marking.upper()).bold = True

        org = doc.add_paragraph()
        org.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = org.add_run(content.org_name)
        run.bold = True

        # реквізит 15 — гриф обмеження доступу (ст.21 З-ну 2657-XII), праворуч угорі
        if content.access_restriction is not None:
            ar = doc.add_paragraph()
            ar.paragraph_format.left_indent = Mm(document.left_indents.restriction_mm)
            ar.add_run(content.access_restriction.heading).bold = True

        # 09 назву виду не зазначають на листах (§4.4)
        if not document.is_letter and content.doc_type.strip():
            dt = doc.add_paragraph()
            dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = dt.add_run(content.doc_type.upper())
            run.font.size = Pt(document.typography.doc_type_size_pt)
            run.bold = True

        # 10 дата + 11 реєстраційний індекс — рядок реквізитів
        meta = doc.add_paragraph()
        meta.add_run(f"{content.date_text}    № {content.reg_index}")

    def _add_approval(self, doc, document: Document, content: DocumentContent) -> None:
        """21 гриф затвердження — праворуч угорі, відступ 100 мм (§7.7)."""
        grant = content.approval
        if grant is None:
            return
        indent = Mm(document.left_indents.approval_mm)

        def _para(text: str, *, bold: bool = False):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = indent
            run = p.add_run(text)
            run.bold = bold
            return p

        _para(grant.heading, bold=True)
        if grant.is_by_document:
            for part in grant.document_reference.split("\n"):
                _para(part)
            return
        if grant.position:
            _para(grant.position)
        if grant.name:
            _para(grant.name)
        if grant.date:
            _para(grant.date)

    def _add_agreements_and_visas(
        self, doc, document: Document, content: DocumentContent
    ) -> None:
        """23 грифи погодження (ПОГОДЖЕНО) та 24 візи — нижче підпису, зліва."""
        for agreement in content.agreements:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("ПОГОДЖЕНО").bold = True
            if agreement.is_by_document:
                for part in agreement.document_reference.split("\n"):
                    doc.add_paragraph().add_run(part)
                continue
            if agreement.position:
                doc.add_paragraph().add_run(agreement.position)
            if agreement.name or agreement.date:
                decode = agreement.name
                if agreement.date:
                    decode = f"{decode}\t\t{agreement.date}".strip()
                doc.add_paragraph().add_run(decode)

        for visa in content.visas:
            doc.add_paragraph()
            doc.add_paragraph().add_run(visa.position)
            decode = visa.name
            if visa.date:
                decode = f"{decode}\t\t{visa.date}".strip()
            doc.add_paragraph().add_run(decode)
            if visa.remark:
                doc.add_paragraph().add_run(f"Зауваження: {visa.remark}")

    def _add_addressees(self, doc, document: Document, content: DocumentContent) -> None:
        if not content.addressees:
            return
        # §7.7: «Адресат» — відступ 90 мм від лівого поля
        for addressee in content.addressees:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(document.left_indents.addressee_mm)
            p.add_run(addressee)

    def _add_title(self, doc, document: Document, content: DocumentContent) -> None:
        if not content.title.strip():
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(content.title).bold = True

    def _add_body(self, doc, document: Document, content: DocumentContent) -> None:
        for para_text in content.body:
            p = doc.add_paragraph()
            # §7.7: абзацний відступ 10 мм
            p.paragraph_format.first_line_indent = Mm(document.left_indents.paragraph_mm)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(para_text)

    def _add_signature(self, doc, document: Document, content: DocumentContent) -> None:
        doc.add_paragraph()
        # підпис (§4.4 реквізит 22):
        #   е-документ із відміткою КЕП → рамка-відмітка по ключу (Art.18/24)
        #   з QR-кодом поряд (§5.10);
        #   інакше → рукописний реквізит «посада + розшифрування».
        if document.is_electronic and content.signatures:
            for mark in content.signatures:
                self._add_e_signature_mark(doc, mark)
                doc.add_paragraph()
            return

        p = doc.add_paragraph()
        # посада ліворуч, розшифрування — у фіксованій колонці на позиції
        # відступу розшифрування (§7.7, 125 мм), щоб імена вирівнювались
        decode_pos = Mm(document.left_indents.signature_decode_mm)
        for i, (position, name) in enumerate(content.paper_signers):
            if i:
                doc.add_paragraph()  # зазор між підписантами (голова/секретар)
                p = doc.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(decode_pos, WD_TAB_ALIGNMENT.LEFT)
            p.add_run(position)
            p.add_run("\t")
            p.add_run(name)

    # --- §4.4(22) ↔ Закон 2155-VIII: відмітка про КЕП ---
    def _add_e_signature_mark(self, doc, mark) -> None:
        """Відмітка про електронний підпис у вигляді рамки з QR поряд.

        Паритет із PdfDocumentWriter: ліворуч — рамка з даними сертифіката
        (підписувач, серійник, видавець, чинність, позначка часу, статус),
        праворуч — QR-код 21×21 мм свого підписувача. Реалізовано як таблиця
        1×2: ліва комірка з рамкою-текстом, права — з картинкою QR.
        """
        lines: list[tuple[str, bool]] = [
            (mark.signature_kind, True),
            (f"Підписувач: {mark.signer}", False),
        ]
        # посада — лише якщо присутня у сертифікаті (сертифікат працівника)
        if mark.signer_position.strip():
            lines.append((f"Посада: {mark.signer_position}", False))
        lines += [
            (f"Сертифікат: {mark.certificate_serial}", False),
            (f"Видавець: {mark.issuer}", False),
            (f"Чинний: {mark.valid_from} – {mark.valid_to}", False),
            (f"Позначка часу: {mark.timestamp}", False),
        ]
        if mark.certificate_valid:
            lines.append(("Статус сертифіката: ЧИННИЙ", True))
        else:
            lines.append(("Статус сертифіката: НЕДІЙСНИЙ (ст.24)", True))

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        mark_cell, qr_cell = table.rows[0].cells
        # §7.6: ширина реквізиту ≤73–95 мм для рамки-відмітки
        mark_cell.width = Mm(95)
        qr_cell.width = Mm(_QR_SIDE_MM + 4)

        # межі лише навколо комірки відмітки (рамка), QR-комірка без меж
        self._set_cell_border(mark_cell)

        first = True
        for text, bold in lines:
            para = mark_cell.paragraphs[0] if first else mark_cell.add_paragraph()
            first = False
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(text)
            run.font.size = Pt(_MARK_FONT_PT)
            run.bold = bold
            if not mark.certificate_valid and text.startswith("Статус"):
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        self._add_qr_image(qr_cell, mark)

    def _add_qr_image(self, cell, mark) -> None:
        """Вставити QR-код 21×21 мм у комірку (кодує дані КЕП за §5.10).

        segno за замовчуванням пише 1-бітний PNG; частина рендерерів .docx
        (LibreOffice, перегляд macOS) показує такий монохром як порожній
        прямокутник. Тож приводимо до RGB перед вставкою, якщо доступний PIL.
        """
        import segno  # локальний імпорт: залежність потрібна лише за наявності QR

        payload = build_signature_qr_payload(mark)
        qr = segno.make(payload, error="m")
        buf = io.BytesIO()
        qr.save(buf, kind="png", scale=10, border=2)
        buf.seek(0)

        try:
            from PIL import Image  # reportlab тягне Pillow; для docx — необов'язково

            rgb = Image.open(buf).convert("RGB")
            buf = io.BytesIO()
            rgb.save(buf, format="PNG")
            buf.seek(0)
        except Exception:  # noqa: BLE001 — без PIL вставляємо як є
            buf.seek(0)

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(buf, width=Mm(_QR_SIDE_MM), height=Mm(_QR_SIDE_MM))

    def _set_cell_border(self, cell) -> None:
        """Намалювати тонку рамку навколо комірки відмітки КЕП."""
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.makeelement(qn("w:tcBorders"), {})
        for edge in ("top", "left", "bottom", "right"):
            elem = borders.makeelement(
                qn(f"w:{edge}"),
                {
                    qn("w:val"): "single",
                    qn("w:sz"): "6",  # 6 восьмих пункта ≈ 0.75 pt
                    qn("w:space"): "0",
                    qn("w:color"): "000000",
                },
            )
            borders.append(elem)
        tc_pr.append(borders)
